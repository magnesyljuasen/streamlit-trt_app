import streamlit as st
from PIL import Image
import requests
from streamlit_lottie import st_lottie
import streamlit_authenticator as stauth
import yaml
from datetime import datetime
import pandas as pd
import numpy as np
from st_aggrid import AgGrid, GridOptionsBuilder
from streamlit_extras.chart_container import chart_container
from trt import trt
import random
#import hydralit_components as hc
#import time
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import io
import datetime
import openai
from io import BytesIO
#from html2image import Html2Image

st.set_page_config(
    page_title="Termisk responstest",
    page_icon="♨️",
    initial_sidebar_state='collapsed'
)

#with hc.HyLoader('', hc.Loaders.standard_loaders,index=[0]):
#    time.sleep(1)

with open("styles/main.css") as f:
    st.markdown("<style>{}</style>".format(f.read()), unsafe_allow_html=True)

with open('src/login/config.yaml') as file:
    config = yaml.load(file, Loader=stauth.SafeLoader)

authenticator = stauth.Authenticate(config['credentials'],config['cookie']['name'],config['cookie']['key'],config['cookie']['expiry_days'])
name, authentication_status, username = authenticator.login('Innlogging', 'main')

if authentication_status == False:
    st.error('Ugyldig brukernavn/passord')
elif authentication_status == None:
    c1, c2, c3 = st.columns(3)
    image = Image.open('src/data/img/logo.png')
    st.image(image)
elif authentication_status:
    with st.sidebar:
        authenticator.logout('Logg ut', 'sidebar')
    #--
    st.button("Refresh")
    st.title("En-to-tre TRT")
    st.markdown("---")
    st.header("1) Inndata")
    project_name = st.selectbox("Velg prosjekt (dette hentes fra frontend, trt-database)", options=["Oslo", "Trondheim", "Bergen"])
    #project_name = st.text_input("Navn på prosjektet")
    well_placement = st.selectbox("Plassering av brønn", options=["Plasser i kart", "Skriv inn koordinater", "Skriv inn adresse"])
    c1, c2 = st.columns(2)
    if well_placement == "Skriv inn koordinater":
        with c1:
            lat = st.number_input("ØV-koordinat")
        with c2:
            long = st.number_input("NS-koordinat")
    well_id_granada = st.text_input("Brønn ID, GRANADA")
    contact_person = st.text_input("Kontaktperson")
    st.markdown("---")
    tab1, tab2, tab3 = st.tabs(["Temperaturprofil", "Brønn og kollektor", "Testdata"])
    with tab1:
        df = pd.DataFrame({
            "Dybde [m]" : list(np.arange(0,310,5).ravel()),
            "Temperatur før [°C]" : list(np.zeros(62).ravel()),
            "Temperatur etter [°C]" : list(np.zeros(62).ravel())
            })
        gb = GridOptionsBuilder.from_dataframe(df)
        gb.configure_default_column(editable=True)
        grid_table = AgGrid(
            df,
            height=400,
            gridOptions=gb.build(),
            fit_columns_on_grid_load=True,
            allow_unsafe_jscode=True,
            key="temperaturprofil",
        )
        grid_table_df = pd.DataFrame(grid_table['data'])
        st.line_chart(data = grid_table_df, y = ["Temperatur før [°C]", "Temperatur etter [°C]"])
    with tab2:
        colletor_type = st.selectbox("Kollektortype", options=["Enkel-U 45 mm", "Enkel-U 40 mm", "Dobbel-U 32 mm"])
        collector_material = st.selectbox("Kollektormateriale", options=["Glatt", "Turbo"])
        collector_length = st.number_input("Kollektorlengde [m]", min_value = 0, value = 300, max_value = 500, step = 10)
        collector_fluid = st.selectbox("Kollektorvæske", options=["HX24", "HX35", "Kilfrost GEO"])
        diameter_casing = st.number_input("Diameter fôringsrør [mm]", value = 139)
        diameter_borehole = st.number_input("Borehullsdiameter [mm]", value = 115)
        groundwater_table = st.number_input("Grunnvannstand [m]")
    with tab3:
        testdata_file = st.file_uploader("Last opp testdata")
        if not testdata_file:
            st.stop()
    st.markdown("---")
    st.header("2) Analyse av termisk responstest")
    df = pd.read_csv(testdata_file, sep = ',', skiprows = 3)
    df = df.iloc[:,[0,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,25,26,28]]
    df = df.rename(columns={
        'Unnamed: 0':'Tidspunkt',
        'Smp.3' :'Panel pipe length',
        'Smp.4' : 'RA_temp(1), Tr_bPump',
        'Smp.5' : 'RA_temp(2), Tr_aPump',
        'Smp.6' : 'RA_temp(3), Theat',
        'Smp.7' : 'RA_temp(4), T_in',
        'Smp.8' : 'RA_temp(5), T_air',
        'Smp.9' : 'Flow',
        'Smp.10' : 'Pump speed',
        'Smp.11' : 'Panel heat',
        'Smp.12' : 'Test_Run',
        'Smp.13' : 'Test_Done',
        'Smp.14' : 'Purge',
        'Smp.15' : 'Purge_Done',
        'Smp.16' : 'Heating',
        'Smp.17' : 'Heating_Done',
        'Smp.18' : 'Cooling',
        'Smp.19' : 'Cooling_Done',
        'Smp.20' : 'Heat_Enabled',
        'Smp.21' : 'Pump_Enabled',
        'Smp.23' : 'Pause_Mode',
        'Smp.24' : 'Mains_Fail',
        'Smp.26' : 'E_Stop',
        })
    matching_rows = df.loc[df['Test_Done'] == -1].index
    test_namelist = []
    test_startrowlist = []
    test_endrowlist = []
    for i in range(1, len(matching_rows)):
        start = matching_rows[i-1] + 10
        stop = matching_rows[i]
        test_df = df.iloc[start:stop]
        if len(test_df) > 1000:
            date_start = test_df['Tidspunkt'][start].split()[0]     
            date_stop = test_df['Tidspunkt'][stop-1].split()[0]
            timestamp_start = f"{date_start.split('-')[2]}/{date_start.split('-')[1]}"
            timestamp_stop = f"{date_stop.split('-')[2]}/{date_stop.split('-')[1]}"
            timestamp_str = f"{timestamp_start} til {timestamp_stop}"
            test_namelist.append(f"{timestamp_str}")
            test_startrowlist.append(start)
            test_endrowlist.append(stop)
    #--
    selected_test = st.selectbox("Velg test", options = test_namelist)
    selected_test_index = test_namelist.index(selected_test)
    selected_df = df.iloc[test_startrowlist[selected_test_index]:test_endrowlist[selected_test_index]]
    with chart_container(selected_df):
        # Create the figure and axes
        fig, ax1 = plt.subplots()

        # Plot the first y-series (Sales) on the left y-axis
        ax1.plot(selected_df.reset_index()["index"], selected_df["RA_temp(2), Tr_aPump"], label='Opp av brønn')
        ax1.plot(selected_df.reset_index()["index"], selected_df["RA_temp(1), Tr_bPump"], label='Ned i brønn')
        ax1.legend(loc='upper left')
        plt.grid(True)
        plt.xlabel("Indeks")
        plt.ylabel("Temperatur [grader]")
        plt.savefig("chart1.png", bbox_inches='tight')
        st.pyplot(plt)
        plt.close()
    with st.expander("Effektiv varmledningsevne"):
        st.write("Matematikk...")
        fig, ax1 = plt.subplots()
        ax1.plot(selected_df["RA_temp(2), Tr_aPump"], selected_df["RA_temp(1), Tr_bPump"])
        plt.grid(True)
        plt.xlabel("Indeks")
        plt.ylabel("Temperatur [grader]")
        plt.savefig("chart2.png", bbox_inches='tight')
        st.pyplot(plt)
    with st.expander("Termisk borehullsmotstand"):
        st.write("Matematikk...")
        st.write("Figur...")
    with st.expander("Temperaturer"):
        st.write("Matematikk...")
        st.write("Figur...")
    #--
    thermal_conductivity = round(random.uniform(2, 5),2)
    borehole_thermal_resistance = 0.08
    undisturbed_groundtemperature = round(random.uniform(6, 10),0)
    
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("Effektiv varmeledningsevne", value = f"{thermal_conductivity} W/(m*K)")
    with c2:
        st.metric("Termisk borehullsmotstand", value = f"{borehole_thermal_resistance} (m*K)/W")
    with c3:
        st.metric("Uforstyrret temperatur", value = f"{undisturbed_groundtemperature} °C")
    st.markdown("---")
    result_text = f"Effektiv varmeledningsevne er målt til å være {thermal_conductivity} W/(m*K). Termisk borehullsmotstand er {borehole_thermal_resistance} (m*K)/W og uforstyrret temperatur er {undisturbed_groundtemperature} °C. Erfaring tilsier at borehullsmotstanden er høyere ved varmeuttak og derfor anbefales det å bruke en borehullsmotstand på {borehole_thermal_resistance + 0.02} (m*K)/W i dimensjoneringen av brønnparken."
    
    document = Document("src/data/docx/Rapportmal.docx")
    styles = document.styles
    style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)

    st.header("3️) Til rapport (og database)")
    #--
    with st.form("Input"):
        forfatter = st.text_input("Forfatter", value="Ola Nordmann")
        oppdragsleder = st.text_input("Oppdragsleder", value="Kari Nordmann")
        oppdragsgiver = st.text_input("Oppdragsgiver", value = "Firma AS")
        oppdragsnummer = st.text_input("Oppdragsnummer", value = "635960-01")
        sted = st.text_input("Sted", value = "Trondheim")
        #--
        depth = st.number_input("Brønndybde [m]", value=300, step=1)
        st.form_submit_button("Gi input")
    #--
    document = Document("src/data/docx/Notatmal.docx")
    styles = document.styles
    style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
    document.paragraphs[0].text = f"Oppdragsgiver: {oppdragsgiver}"
    document.paragraphs[1].text = f"Oppdragsnavn: Termisk responstest - {sted}"
    document.paragraphs[2].text = f"Oppdragsnummer: {oppdragsnummer} - {sted}"
    document.paragraphs[3].text = f"Utarbeidet av: {forfatter}"
    document.paragraphs[4].text = f"Oppdragsleder: {oppdragsleder}"
    document.paragraphs[5].text = f"Dato: {datetime.date.today()}"
    document.paragraphs[6].text = f"Tilgjenglighet: Åpen"

    document.paragraphs[7].text = f"Termisk responstest - {sted}"
    report_text_1 = f"""Asplan Viak har på oppdrag for {oppdragsgiver} analysert og rapportert resultatene fra en termisk responstest. Responstesten er utført i et {depth} m dypt borehull ved {sted}. Formålet med den termiske responstesten er å finne stedsspesifikke parametere som grunnlag for riktig dimensjonering av et grunnvarmeanlegg. For å dimensjonere grunnvarmeanlegget med utgangspunkt i resultatene fra testen må det brukes egnet programvare, f.eks. EED (Earth Energy Designer), som ivaretar varmetransporten i berggrunnen og interaksjonen mellom energibrønnene. Rapporten beskriver hvilke undersøkelser og beregninger som er gjort, og resultatene fra disse. """
    #-- AI
    #st.header("Tekst med Open AI")
    #openai.api_key = st.secrets["openai_apikey"]
    # Set up the model and prompt
    #model_engine = "text-davinci-003"
    #prompt = st.text_input("Hva skal jeg skrive?", placeholder=f"Du er en geolog. Skriv et kort avsnitt om geologien i {sted}")
    #temperature = st.slider("Kreativitet", min_value=0.0, value=0.5, max_value=1.0, step=0.1)
    #with st.spinner("Skriver..."):
    #    if prompt:
    #        st.header("Respons")
            # Generate a response
    #        completion = openai.Completion.create(
    #            engine=model_engine,
    #            prompt=prompt,
    #            max_tokens=4000,
    #            n=1,
    #            stop=None,
    #            temperature=temperature,
    #        )

    #        for i in range(0, len(completion.choices)):
    #            st.markdown("---")
    #            geology_text = (completion.choices[i].text)
    #--
    
    document.add_heading("Innledning", 1)
    document.add_paragraph(report_text_1)
    
    document.add_heading("Områdebeskrivelse")
    document.add_paragraph("Geologi...")
        
        
    document.add_heading("Analyse av termisk responstest")
    #p = document.add_paragraph()
    #r = p.add_run()
    #r.add_text(result_text)
    #r.add_picture("chart1.png")
    #r.add_picture("chart2.png")
    document.add_paragraph(result_text)
    document.add_picture("chart1.png")
    document.add_picture("chart2.png")
        
        #--
        #c1, c2 = st.columns(2)
        #with st.form("Inndata"):
        #    with c1:
        #        TITTEL = st.text_input("Tittel", value = f"Termisk responstest - {project_name}")
        #        OPPDRAGSNAVN = st.text_input("Oppdragsnavn", value = f"Termisk responstest - {project_name}")
        #        OPPDRAGSGIVER = st.text_input("Oppdragsgiver", placeholder= "Båsum Boring")
        #        OPPDRAGSNUMMER = st.text_input("Oppdragsnummer", placeholder= "630962-01")
        #    with c2:
        #        FORFATTER = st.selectbox("Utarbeidet av", options = ["Johanne Strålberg", "Sofie Hartvigsen", "Magne Syljuåsen", "Henrik Holmberg", "Randi Kalskin Ramstad"])
        #        OPPDRAGSLEDER = st.selectbox("Oppdragsleder", options = ["Johanne Strålberg", "Sofie Hartvigsen", "Magne Syljuåsen", "Henrik Holmberg", "Randi Kalskin Ramstad"])
        #        KVALITETSSIKRER = st.selectbox("Kvalitetssikrer", options = ["Magne Syljuåsen", "Henrik Holmberg", "Randi Kalskin Ramstad", "Johanne Strålberg", "Sofie Hartvigsen"])
        #    st.form_submit_button("Gi input")

    # document.paragraphs[1].text = f"Oppdragsgiver:      {OPPDRAGSGIVER}"
    # document.paragraphs[2].text = f"Oppdragsgiver:      {OPPDRAGSGIVER}"
    # document.paragraphs[3].text = f"Tittel på rapport:      {OPPDRAGSNAVN}"
    # document.paragraphs[4].text = f"Tittel på rapport:      {OPPDRAGSNAVN}"
    # st.write("---")

    #--
    bio = io.BytesIO()
    document.save(bio)
    if document:
        st.download_button(
            label="Last ned rapport!",
            data=bio.getvalue(),
            file_name="Rapport.docx",
            mime="docx")


        
        







